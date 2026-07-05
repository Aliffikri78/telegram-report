#!/usr/bin/env python3
import os, sys, re
from pathlib import Path
from typing import List, Tuple
from PIL import Image
import cv2, numpy as np
from tqdm import tqdm, trange
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

def input_default(prompt: str, default: str) -> str:
    v = input(f"{prompt} [{default}]: ").strip()
    return v if v else default

def list_dirs(p: Path) -> List[Path]:
    return [d for d in sorted(p.iterdir()) if d.is_dir()]

def is_month_name(name: str) -> bool:
    return re.match(r"^\d{4}-\d{2}$", name) is not None

def load_images(folder: Path) -> List[Path]:
    exts = {".jpg",".jpeg",".png",".bmp",".tif",".tiff",".webp"}
    if not folder.exists(): return []
    return [p for p in sorted(folder.iterdir()) if p.suffix.lower() in exts]

def crop_to_4x3(im: Image.Image) -> Image.Image:
    w,h = im.size; target=4/3; cur=w/h
    if abs(cur-target) < 0.005: return im
    if cur > target:
        new_w = int(h*target); x0=(w-new_w)//2; return im.crop((x0,0,x0+new_w,h))
    new_h = int(w/target); y0=(h-new_h)//2; return im.crop((0,y0,w,y0+new_h))

def orb_score(a_path: Path, b_path: Path) -> float:
    try:
        a = cv2.imdecode(np.fromfile(str(a_path), dtype=np.uint8), cv2.IMREAD_GRAYSCALE)
        b = cv2.imdecode(np.fromfile(str(b_path), dtype=np.uint8), cv2.IMREAD_GRAYSCALE)
        if a is None or b is None: return 0.0
        orb = cv2.ORB_create(nfeatures=800)
        ka, da = orb.detectAndCompute(a, None)
        kb, db = orb.detectAndCompute(b, None)
        if da is None or db is None: return 0.0
        bf = cv2.BFMatcher(cv2.NORM_HAMMING, crossCheck=True)
        matches = bf.match(da, db)
        if not matches: return 0.0
        dists=[m.distance for m in matches]
        avg = sum(dists)/len(dists)
        return max(0.0, 1.0 - avg/128.0)
    except Exception:
        return 0.0

def best_pairs(befores: List[Path], afters: List[Path], threshold: float=0.70) -> List[Tuple[Path,Path,float]]:
    pairs=[]; used=set()
    print(f"\n[Matching] {len(befores)} before vs {len(afters)} after (threshold={threshold:.2f})")
    for bi, b in enumerate(tqdm(befores, desc="Pairing", unit="img")):
        best_j, best_s=None, -1.0
        for j, a in enumerate(afters):
            if j in used: continue
            s = orb_score(b, a)
            if s > best_s:
                best_s, best_j = s, j
        if best_j is not None and best_s >= threshold:
            used.add(best_j); pairs.append((b, afters[best_j], best_s))
        else:
            j = bi if (bi < len(afters) and bi not in used) else None
            if j is not None:
                used.add(j); pairs.append((b, afters[j], -1.0))
    return pairs

def set_table_borders_none(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tbl_pr.append(borders)

def add_header(doc: Document, company: str, zone: str):
    p = doc.add_paragraph(); r=p.add_run(company.upper()); r.bold=True; r.font.size=Pt(14)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph(zone.upper()); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].font.size=Pt(12)

def add_title_left(doc: Document, title: str):
    p = doc.add_paragraph(); r=p.add_run(title); r.bold=True; r.font.size=Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_pair_table(doc: Document, before_path: Path, after_path: Path):
    sec = doc.sections[-1]
    gutter = Inches(0.3)
    max_w = (sec.page_width - sec.left_margin - sec.right_margin - gutter) / 2
    tbl = doc.add_table(rows=2, cols=2)
    set_table_borders_none(tbl)
    try:
        im = Image.open(before_path).convert("RGB"); im = crop_to_4x3(im)
        tmp = before_path.parent / f"~tmp_{before_path.stem}_b.jpg"
        im.save(tmp, quality=90)
        cell = tbl.cell(0,0).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.add_run().add_picture(str(tmp), width=max_w)
        tmp.unlink(missing_ok=True)
    except Exception:
        tbl.cell(0,0).text="(gambar gagal)"
    try:
        im = Image.open(after_path).convert("RGB"); im = crop_to_4x3(im)
        tmp = after_path.parent / f"~tmp_{after_path.stem}_a.jpg"
        im.save(tmp, quality=90)
        cell = tbl.cell(0,1).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cell.add_run().add_picture(str(tmp), width=max_w)
        tmp.unlink(missing_ok=True)
    except Exception:
        tbl.cell(0,1).text="(gambar gagal)"
    pL = tbl.cell(1,0).paragraphs[0]; pL.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rL = pL.add_run("SEBELUM"); rL.bold = True; rL.font.size = Pt(11)
    pR = tbl.cell(1,1).paragraphs[0]; pR.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rR = pR.add_run("SELEPAS"); rR.bold = True; rR.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def generate_report(input_root: Path, out_path: Path, company: str, zone: str, title: str, threshold: float=0.70):
    before_dir = input_root / "before"; after_dir = input_root / "after"
    befores = load_images(before_dir); afters = load_images(after_dir)
    if not befores or not afters:
        print("ERROR: No images in before/ or after/"); sys.exit(2)
    pairs = best_pairs(befores, afters, threshold)
    doc = Document()
    add_header(doc, company, zone)
    add_title_left(doc, title)
    per_page = 4
    for i in range(min(per_page, len(pairs))):
        add_pair_table(doc, pairs[i][0], pairs[i][1])
    if len(pairs) > per_page:
        for idx in trange(per_page, len(pairs), per_page, desc="Pages", unit="page"):
            doc.add_page_break()
            add_header(doc, company, zone)
            for k in range(per_page):
                j = idx + k
                if j >= len(pairs): break
                add_pair_table(doc, pairs[j][0], pairs[j][1])
    doc.save(out_path)
    print("Saved:", out_path)

def main():
    print("=== Telegram Report Builder (TABLE) ===")
    default_root = "/data/photos" if os.getenv("IN_DOCKER")=="1" else ("/Volumes/NAS/telegram_report_automation/photos" if sys.platform=="darwin" else "/data/photos")
    root = Path(input_default("Photos ROOT folder", default_root)).expanduser().resolve()
    if not root.exists(): print("ERROR: root not found:", root); sys.exit(2)
    months = [d for d in list_dirs(root) if is_month_name(d.name)]
    if not months: print("No month folders (YYYY-MM) in", root); sys.exit(2)
    print("\nSelect month:"); 
    for i,m in enumerate(months,1): print(f"  {i}. {m.name}")
    mi = int(input_default("Enter number", "1")) - 1; month_dir = months[max(0,min(mi,len(months)-1))]
    sites = [d for d in list_dirs(month_dir) if (d/\"grass_cutting\").exists() or (d/\"drainage_cleaning\").exists()]
    if not sites: print("No site folders under", month_dir); sys.exit(2)
    print("\nSelect site:"); 
    for i,s in enumerate(sites,1): print(f"  {i}. {s.name}")
    si = int(input_default("Enter number", "1")) - 1; site_dir = sites[max(0,min(si,len(sites)-1))]
    print("\nSelect task:")
    options=[]
    if (site_dir/\"grass_cutting\").exists(): options.append((\"grass_cutting\",\"1\"))
    if (site_dir/\"drainage_cleaning\").exists(): options.append((\"drainage_cleaning\",\"2\"))
    if not options: print("No tasks found in", site_dir); sys.exit(2)
    for t,n in options: print(f"  {n}. {t}")
    choice = input_default("Enter number", options[0][1])
    task = "grass_cutting" if choice=="1" else "drainage_cleaning"
    default_zone = f"{site_dir.name} ZONE"
    company = input_default("Company", "HW UNGGUL (901587-V)")
    zone    = input_default("Zone", default_zone)
    title   = input_default("Section title", "1. Grass Cutting" if task=="grass_cutting" else "2. Drainage Cleaning")
    thresh  = float(input_default("Match threshold (0.50-0.95)", "0.70"))
    outdir  = Path(input_default("Output folder", "/data/reports" if os.getenv("IN_DOCKER")=="1" else ("/Volumes/NAS/telegram_report_automation/reports" if sys.platform=='darwin' else "/data/reports"))).expanduser()
    outdir.mkdir(parents=True, exist_ok=True)
    input_root = site_dir / task
    month_name = month_dir.name
    out_name = f"{month_name}_{site_dir.name}_{'grass_cutting' if task=='grass_cutting' else 'drainage'}.docx"
    out_path = outdir / out_name
    generate_report(input_root, out_path, company, zone, title, thresh)

if __name__ == "__main__":
    main()
