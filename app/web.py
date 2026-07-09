#!/usr/bin/env python3
import csv, json, os, re, secrets, subprocess, threading, time, uuid
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime, timedelta
from functools import wraps
from pathlib import Path
from typing import List, Dict
from urllib import parse, request as urlrequest
from flask import Flask, render_template, request, Response, jsonify, send_file, make_response, redirect, session, url_for

from PIL import Image
import cv2
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

try:
    from vision import ai_engine
except ImportError:
    try:
        from app.vision import ai_engine
    except ImportError:
        ai_engine = None

try:
    import importlib.util

    from core.logger import logger
    from database.db import connection
    from jobs.manager import JobManager
    from projects.manager import ProjectManager
    from storage.storage import Storage
    from vision.analyze import analyze_folder
    from vision.cache import VisionCache
    from vision.matcher import match_pairs
    from worker_uploads.manager import WorkerUploadManager

    queue_manager_path = Path(__file__).resolve().parent / "queue" / "manager.py"
    spec = importlib.util.spec_from_file_location("local_queue_manager", queue_manager_path)
    queue_manager_module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(queue_manager_module)
    QueueManager = queue_manager_module.QueueManager
except ImportError:
    from app.core.logger import logger
    from app.database.db import connection
    from app.jobs.manager import JobManager
    from app.projects.manager import ProjectManager
    from app.queue.manager import QueueManager
    from app.storage.storage import Storage
    from app.vision.analyze import analyze_folder
    from app.vision.cache import VisionCache
    from app.vision.matcher import match_pairs
    from app.worker_uploads.manager import WorkerUploadManager

storage = Storage()
project_manager = ProjectManager()
DATA_ROOT = storage.photos_root
REPORT_ROOT = storage.ensure_dir(storage.reports_root)

def list_dirs(p: Path):
    return storage.list_dirs(p)

def is_month_name(name: str) -> bool:
    import re
    return re.match(r"^\d{4}-\d{2}$", name) is not None

def scan_months_sites() -> Dict[str, list]:
    out = {}
    if not DATA_ROOT.exists(): return out
    configured = {task["slug"] for task in project_manager.list_tasks()}
    legacy = {"grass_cutting", "drainage_cleaning"}
    task_slugs = configured | legacy
    for m in list_dirs(DATA_ROOT):
        if not is_month_name(m.name): continue
        sites = []
        for s in list_dirs(m):
            if any((s / task_slug).exists() for task_slug in task_slugs):
                sites.append(s.name)
        if sites:
            out[m.name] = sites
    return out

def report_tasks():
    tasks = project_manager.list_tasks()
    by_slug = {task["slug"]: task for task in tasks}
    ordered = []
    for slug in ("grass_cutting", "drainage_cleaning"):
        if slug in by_slug:
            ordered.append(by_slug.pop(slug))
    ordered.extend(by_slug.values())
    return ordered

def report_defaults():
    tasks = report_tasks()
    default_task = next(
        (task for task in tasks if task["slug"] == "grass_cutting"),
        tasks[0] if tasks else {"slug": "", "title": "1. Grass Cutting", "name": "Grass Cutting"},
    )
    return {
        "company": os.getenv("COMPANY", "HW UNGGUL (901587-V)"),
        "zone": "ZONE",
        "title": default_task.get("title") or default_task.get("name") or "1. Grass Cutting",
        "threshold": os.getenv("THRESHOLD", "0.70"),
        "task": default_task["slug"],
    }

IMG_H_CM = float(os.getenv("IMG_H_CM", "4.3"))

class CancelledJob(Exception):
    pass

def ai_label(confidence: float) -> str:
    if confidence >= 0.70:
        return "High"
    if confidence >= 0.40:
        return "Review"
    return "Low"

def summarize_ai_results(results):
    if not results:
        return {
            "total": 0,
            "high_count": 0,
            "review_count": 0,
            "low_count": 0,
            "average_ai_confidence": 0.0,
            "lowest_ai_confidence": 0.0,
            "highest_ai_confidence": 0.0,
            "average_processing_time_ms": 0.0,
        }

    confidences = [float(item.get("confidence") or 0.0) for item in results]
    times = [float(item.get("processing_time_ms") or 0.0) for item in results]
    return {
        "total": len(results),
        "high_count": sum(1 for item in results if item.get("label") == "High"),
        "review_count": sum(1 for item in results if item.get("label") == "Review"),
        "low_count": sum(1 for item in results if item.get("label") == "Low"),
        "average_ai_confidence": round(sum(confidences) / len(confidences), 4),
        "lowest_ai_confidence": round(min(confidences), 4),
        "highest_ai_confidence": round(max(confidences), 4),
        "average_processing_time_ms": round(sum(times) / len(times), 2),
    }

def calibration_enabled():
    return os.getenv("AI_CALIBRATION_ENABLED", "").lower() in {"1", "true", "yes", "on"}

def calibration_csv_path():
    return Path(os.getenv("AI_CALIBRATION_CSV", "/data/reports_dev/ai_calibration.csv"))

def ai_recovery_enabled():
    return os.getenv("AI_RECOVERY_ENABLED", "").lower() in {"1", "true", "yes", "on"}

def ai_recovery_threshold():
    try:
        return float(os.getenv("AI_RECOVERY_THRESHOLD", "0.55"))
    except ValueError:
        return 0.55

def ai_recovery_top_candidates():
    try:
        return max(1, int(os.getenv("AI_RECOVERY_TOP_CANDIDATES", "30")))
    except ValueError:
        return 30

def bool_from_value(value):
    return str(value or "").lower() in {"1", "true", "yes", "on"}

def timestamp_score(path: Path):
    digits = re.findall(r"\d+", Path(path).stem)
    values = []
    for item in digits:
        if len(item) >= 13:
            values.append(int(item[:13]))
        elif len(item) >= 12:
            values.append(int(item[:12]))
        elif len(item) >= 8:
            values.append(int(item[:8]))
    return max(values) if values else None

def ai_status_payload(health):
    if isinstance(health, dict) and isinstance(health.get("ai"), dict):
        return health["ai"]
    return health if isinstance(health, dict) else {}

def ai_recovery_backend_status(progress):
    health = progress.get("ai_service_health")
    if not health and ai_engine is not None and hasattr(ai_engine, "health"):
        try:
            health = ai_engine.health()
            progress["ai_service_health"] = health
        except Exception as exc:
            health = {"ok": False, "error": str(exc)}
            progress["ai_service_health"] = health
    ai_status = ai_status_payload(health)
    service_ok = bool(health.get("ok")) if isinstance(health, dict) else False
    cuda_available = ai_status.get("cuda_available") if isinstance(ai_status, dict) else None
    gpu_name = (ai_status.get("gpu_name") or ai_status.get("device")) if isinstance(ai_status, dict) else None
    fallback_reason = None
    if not service_ok:
        backend = "cpu"
        fallback_reason = (health or {}).get("error") if isinstance(health, dict) else "AI service unavailable"
        fallback_reason = fallback_reason or "AI service unhealthy; AI Recovery skipped"
    elif cuda_available is True:
        backend = "cuda"
    elif cuda_available is False:
        backend = "cpu"
        fallback_reason = "CUDA unavailable; AI Recovery using CPU service backend"
    else:
        backend = "cpu"
        fallback_reason = "CUDA status unknown; AI Recovery using CPU service backend"
    return {
        "backend": backend,
        "service_ok": service_ok,
        "cuda_available": cuda_available if isinstance(cuda_available, bool) else None,
        "gpu_name": gpu_name,
        "fallback_reason": fallback_reason,
    }

def filtered_recovery_candidates(before_path: Path, remaining_after, top_n: int, deep_recovery: bool):
    if deep_recovery:
        return list(remaining_after)
    before_ts = timestamp_score(before_path)
    ranked = []
    for idx, after_path in enumerate(remaining_after):
        after_ts = timestamp_score(after_path)
        if before_ts is not None and after_ts is not None:
            distance = abs(before_ts - after_ts)
        else:
            distance = 10**18
        ranked.append((distance, Path(after_path).name, idx, after_path))
    ranked.sort()
    return [item[3] for item in ranked[:top_n]]

def safe_float(value, default=0.0):
    try:
        return float(value)
    except (TypeError, ValueError):
        return default

def check_ai_service(progress):
    if ai_engine is None:
        health = {"ok": False, "error": "AI engine unavailable"}
    elif hasattr(ai_engine, "health"):
        try:
            health = ai_engine.health()
        except Exception as exc:
            health = {"ok": False, "error": f"AI service health check failed: {exc}"}
    else:
        health = {"ok": False, "error": "AI service health check unavailable"}
    progress["ai_service_health"] = health
    return bool(health.get("ok"))

def append_ai_calibration_rows(report_name, results):
    if not calibration_enabled() or not results:
        return

    path = calibration_csv_path()
    path.parent.mkdir(parents=True, exist_ok=True)
    fields = [
        "created_at",
        "report_name",
        "before_image",
        "after_image",
        "cpu_confidence",
        "ai_confidence",
        "matches",
        "keypoints_before",
        "keypoints_after",
        "processing_time_ms",
        "final_label",
        "error",
    ]

    exists = path.exists()
    with open(path, "a", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields)
        if not exists:
            writer.writeheader()

        created_at = datetime.utcnow().isoformat(timespec="seconds") + "Z"
        for item in results:
            writer.writerow(
                {
                    "created_at": created_at,
                    "report_name": report_name,
                    "before_image": item.get("before"),
                    "after_image": item.get("after"),
                    "cpu_confidence": item.get("cpu_score"),
                    "ai_confidence": item.get("confidence"),
                    "matches": item.get("matches"),
                    "keypoints_before": item.get("keypoints_before"),
                    "keypoints_after": item.get("keypoints_after"),
                    "processing_time_ms": item.get("processing_time_ms"),
                    "final_label": item.get("label"),
                    "error": item.get("error"),
                }
            )

def pair_paths(pair):
    if isinstance(pair, dict):
        return pair.get("before"), pair.get("after"), pair.get("score") or pair.get("final_score")
    if isinstance(pair, (list, tuple)) and len(pair) >= 2:
        score = pair[2] if len(pair) > 2 else None
        return pair[0], pair[1], score
    return None, None, None

def run_ai_pair_review(pairs, progress: dict, should_cancel=None):
    should_cancel = should_cancel or (lambda: False)
    results = []
    progress["ai_results"] = results
    progress["ai_error"] = None

    if ai_engine is None:
        progress["ai_state"] = "error"
        progress["ai_error"] = "AI engine unavailable"
        return results

    progress["ai_state"] = "running"
    progress["ai_total"] = len(pairs)
    progress["ai_done"] = 0

    for idx, pair in enumerate(pairs):
        if should_cancel():
            raise CancelledJob("Cancelled by user")

        before_path, after_path, cpu_score = pair_paths(pair)
        if not before_path or not after_path:
            result = {
                "confidence": 0.0,
                "matches": 0,
                "keypoints_before": 0,
                "keypoints_after": 0,
                "processing_time_ms": 0.0,
                "error": "Invalid pair data",
            }
        else:
            try:
                result = ai_engine.match(str(before_path), str(after_path))
            except Exception as exc:
                result = {
                    "confidence": 0.0,
                    "matches": 0,
                    "keypoints_before": 0,
                    "keypoints_after": 0,
                    "processing_time_ms": 0.0,
                    "error": str(exc),
                }

        confidence = float(result.get("confidence") or 0.0)
        result.update({
            "before": str(before_path) if before_path else None,
            "after": str(after_path) if after_path else None,
            "cpu_score": cpu_score,
            "label": ai_label(confidence),
        })
        results.append(result)
        progress["ai_done"] = idx + 1
        progress["ai_results"] = results

    progress["ai_state"] = "done"
    progress["ai_summary"] = summarize_ai_results(results)
    return results

def run_ai_recovery(befores, afters, pairs, progress: dict, should_cancel=None, deep_recovery=False):
    threshold = ai_recovery_threshold()
    stats = {
        "enabled": ai_recovery_enabled(),
        "threshold": threshold,
        "attempted": 0,
        "recovered": 0,
        "failed": 0,
        "comparisons": 0,
        "planned_comparisons": 0,
        "full_comparisons": 0,
        "candidate_pairs_after_filtering": 0,
        "total_fallback_pairs": 0,
        "top_candidates": ai_recovery_top_candidates(),
        "deep_recovery": bool(deep_recovery),
        "backend": "unknown",
        "gpu_name": None,
        "cuda_available": None,
        "fallback_reason": None,
        "service_ok": None,
        "state": "disabled",
        "current_before": None,
        "current_after": None,
        "current_pair": None,
        "eta_seconds": None,
        "started_at": None,
        "completed_at": None,
        "passes": [],
        "results": [],
        "failed_candidates": [],
        "errors": [],
    }
    progress["ai_recovery"] = stats

    if not stats["enabled"]:
        return pairs
    if ai_engine is None:
        stats["state"] = "error"
        stats["completed_at"] = datetime.utcnow().isoformat(timespec="seconds") + "Z"
        stats["errors"].append("AI engine unavailable")
        stats["failed_candidates"].append({
            "before": None,
            "best_after": None,
            "best_confidence": 0.0,
            "threshold": threshold,
            "reason": "ai_engine_unavailable",
        })
        progress["ai_recovery"] = stats
        return pairs

    should_cancel = should_cancel or (lambda: False)
    used_after = set()
    for pair in pairs:
        if isinstance(pair, (list, tuple)) and len(pair) > 1:
            used_after.add(str(pair[1]))
        else:
            _, after_path, _ = pair_paths(pair)
            if after_path:
                used_after.add(str(after_path))

    remaining_after = [path for path in afters if str(path) not in used_after]
    unmatched_before = [
        Path(item["before"])
        for item in progress.get("unmatched_debug", [])
        if item.get("before")
    ]

    backend_status = progress.get("ai_recovery_backend_decision") or ai_recovery_backend_status(progress)
    stats.update(backend_status)
    if not backend_status.get("service_ok"):
        stats["state"] = "skipped"
        stats["completed_at"] = datetime.utcnow().isoformat(timespec="seconds") + "Z"
        stats["eta_seconds"] = 0
        stats["errors"].append(stats.get("fallback_reason") or "AI service unhealthy; AI Recovery skipped")
        progress["ai_recovery"] = stats
        return pairs

    stats["total_fallback_pairs"] = len(unmatched_before)
    stats["full_comparisons"] = len(unmatched_before) * len(remaining_after)

    print(
        "[ai_recovery] fallback_pairs=%s full_pairs=%s backend=%s cuda_available=%s gpu=%s deep=%s top_candidates=%s" % (
            stats["total_fallback_pairs"],
            stats["full_comparisons"],
            stats["backend"],
            stats["cuda_available"],
            stats["gpu_name"],
            stats["deep_recovery"],
            stats["top_candidates"],
        ),
        flush=True,
    )

    recovery_started = time.monotonic()
    last_progress_update = 0.0
    stats["state"] = "running"
    stats["started_at"] = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    progress["ai_recovery"] = stats

    def update_eta():
        elapsed = max(0.0, time.monotonic() - recovery_started)
        completed = int(stats.get("comparisons") or 0)
        total = int(stats.get("planned_comparisons") or 0)
        remaining = max(0, total - completed)
        if completed > 0 and elapsed > 0 and remaining > 0:
            stats["eta_seconds"] = int(round(remaining / (completed / elapsed)))
        else:
            stats["eta_seconds"] = 0 if remaining == 0 else None

    def publish_recovery_progress(force=False):
        nonlocal last_progress_update
        update_eta()
        now = time.monotonic()
        comparisons = int(stats.get("comparisons") or 0)
        if force or comparisons % 10 == 0 or now - last_progress_update >= 0.5:
            progress["ai_recovery"] = stats
            last_progress_update = now

    def failure_record(before_path, best_after, best_confidence, reason):
        return {
            "before": str(before_path) if before_path else None,
            "best_after": str(best_after) if best_after else None,
            "best_confidence": round(best_confidence, 4) if best_confidence >= 0 else 0.0,
            "threshold": threshold,
            "reason": reason,
        }

    def run_recovery_pass(pass_number, before_list, top_n, pass_deep):
        candidate_plan = {
            before_path: filtered_recovery_candidates(before_path, remaining_after, top_n, pass_deep)
            for before_path in before_list
        }
        pass_stats = {
            "pass": pass_number,
            "deep_recovery": bool(pass_deep),
            "top_candidates": top_n,
            "attempted": 0,
            "candidate_pairs": sum(len(items) for items in candidate_plan.values()),
            "comparisons": 0,
            "recovered": 0,
            "failed": 0,
        }
        stats["passes"].append(pass_stats)
        stats["candidate_pairs_after_filtering"] += pass_stats["candidate_pairs"]
        stats["planned_comparisons"] += pass_stats["candidate_pairs"]
        next_unmatched = []
        pass_failures = []

        for before_path in before_list:
            if should_cancel():
                raise CancelledJob("Cancelled by user")
            if not remaining_after:
                pass_stats["failed"] += 1
                next_unmatched.append(before_path)
                pass_failures.append(failure_record(before_path, None, 0.0, "no_remaining_after_images"))
                publish_recovery_progress(force=True)
                continue

            pass_stats["attempted"] += 1
            stats["attempted"] += 1
            best_after = None
            best_result = None
            best_confidence = -1.0

            for after_path in list(candidate_plan.get(before_path, [])):
                if after_path not in remaining_after:
                    continue
                if should_cancel():
                    raise CancelledJob("Cancelled by user")
                stats["current_before"] = Path(before_path).name
                stats["current_after"] = Path(after_path).name
                stats["current_pair"] = {
                    "before": Path(before_path).name,
                    "after": Path(after_path).name,
                }
                stats["comparisons"] += 1
                pass_stats["comparisons"] += 1
                publish_recovery_progress()
                try:
                    result = ai_engine.match(str(before_path), str(after_path))
                except Exception as exc:
                    stats["errors"].append({
                        "before": str(before_path),
                        "after": str(after_path),
                        "error": str(exc),
                    })
                    publish_recovery_progress()
                    continue

                if result.get("error"):
                    stats["errors"].append({
                        "before": str(before_path),
                        "after": str(after_path),
                        "error": result.get("error"),
                    })
                    publish_recovery_progress()
                    continue

                confidence = safe_float(result.get("confidence"), 0.0)
                if confidence > best_confidence:
                    best_confidence = confidence
                    best_after = after_path
                    best_result = result
                publish_recovery_progress()

            if best_after and best_result and best_confidence >= threshold:
                recovered = {
                    **best_result,
                    "before": str(before_path),
                    "after": str(best_after),
                    "confidence": best_confidence,
                    "label": "ai_recovered",
                    "recovery_pass": pass_number,
                }
                stats["recovered"] += 1
                pass_stats["recovered"] += 1
                stats["results"].append(recovered)
                remaining_after.remove(best_after)
                pairs.append((before_path, best_after, best_confidence))
            else:
                pass_stats["failed"] += 1
                next_unmatched.append(before_path)
                pass_failures.append(failure_record(
                    before_path,
                    best_after,
                    best_confidence,
                    "below_threshold" if best_after else "no_valid_ai_candidate",
                ))
            publish_recovery_progress(force=True)

        print(
            "[ai_recovery] pass=%s candidate_pairs=%s comparisons=%s recovered=%s failed=%s deep=%s top_candidates=%s" % (
                pass_number,
                pass_stats["candidate_pairs"],
                pass_stats["comparisons"],
                pass_stats["recovered"],
                pass_stats["failed"],
                pass_stats["deep_recovery"],
                pass_stats["top_candidates"],
            ),
            flush=True,
        )
        return next_unmatched, pass_failures

    remaining_before, final_failures = run_recovery_pass(
        1,
        unmatched_before,
        stats["top_candidates"],
        stats["deep_recovery"],
    )

    pass1 = stats["passes"][0] if stats["passes"] else {}
    pass1_attempted = int(pass1.get("attempted") or 0)
    pass1_recovered = int(pass1.get("recovered") or 0)
    recovery_rate = (pass1_recovered / pass1_attempted) if pass1_attempted else 0.0
    remaining_fraction = (len(remaining_before) / len(unmatched_before)) if unmatched_before else 0.0
    should_escalate = (
        bool(remaining_before)
        and not stats["deep_recovery"]
        and (recovery_rate < 0.50 or remaining_fraction > 0.25)
    )

    if should_escalate:
        widened_top = max(len(remaining_after), stats["top_candidates"] * 3)
        print(
            "[ai_recovery] escalating pass=2 reason=recovery_rate:%0.3f remaining_fraction:%0.3f remaining_before=%s remaining_after=%s" % (
                recovery_rate,
                remaining_fraction,
                len(remaining_before),
                len(remaining_after),
            ),
            flush=True,
        )
        remaining_before, final_failures = run_recovery_pass(2, remaining_before, widened_top, True)

    stats["failed"] = len(remaining_before)
    stats["failed_candidates"] = final_failures

    if stats["recovered"]:
        progress["matched"] = int(progress.get("matched", 0)) + stats["recovered"]
        progress["unmatched"] = max(0, int(progress.get("unmatched", 0)) - stats["recovered"])

    stats["state"] = "done" if stats["state"] == "running" else stats["state"]
    stats["current_before"] = None
    stats["current_after"] = None
    stats["current_pair"] = None
    stats["eta_seconds"] = 0
    stats["completed_at"] = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    print(
        "[ai_recovery] final recovered=%s failed=%s comparisons=%s candidate_pairs=%s" % (
            stats["recovered"],
            stats["failed"],
            stats["comparisons"],
            stats["candidate_pairs_after_filtering"],
        ),
        flush=True,
    )
    publish_recovery_progress(force=True)
    return pairs

def notify_telegram(message: str) -> None:
    token = os.getenv("TG_BOT_TOKEN")
    chat_id = os.getenv("TG_CHAT_ID")
    if not token or not chat_id:
        return

    def send():
        try:
            data = parse.urlencode({"chat_id": chat_id, "text": message}).encode("utf-8")
            req = urlrequest.Request(f"https://api.telegram.org/bot{token}/sendMessage", data=data)
            urlrequest.urlopen(req, timeout=5).close()
        except Exception:
            pass

    threading.Thread(target=send, daemon=True).start()

def detect_gpu_status() -> Dict:
    status = {
        "nvidia_detected": False,
        "gpu_name": None,
        "opencv_cuda_available": False,
        "opencv_cuda_device_count": 0,
        "opencv_cuda_build": False,
        "error": None,
    }
    errors = []

    try:
        result = subprocess.run(
            ["nvidia-smi", "--query-gpu=name", "--format=csv,noheader"],
            check=False,
            capture_output=True,
            text=True,
            timeout=5,
        )
        if result.returncode == 0 and result.stdout.strip():
            names = [line.strip() for line in result.stdout.splitlines() if line.strip()]
            status["nvidia_detected"] = bool(names)
            status["gpu_name"] = names[0] if names else None
        elif result.stderr.strip():
            errors.append(result.stderr.strip())
    except Exception as exc:
        errors.append(str(exc))

    try:
        build_info = cv2.getBuildInformation()
        status["opencv_cuda_build"] = "NVIDIA CUDA" in build_info and "YES" in build_info.split("NVIDIA CUDA", 1)[1][:80]
    except Exception as exc:
        errors.append(f"OpenCV build info unavailable: {exc}")

    try:
        if hasattr(cv2, "cuda"):
            count = cv2.cuda.getCudaEnabledDeviceCount()
            status["opencv_cuda_device_count"] = int(count)
            status["opencv_cuda_available"] = count > 0
    except Exception as exc:
        errors.append(f"OpenCV CUDA unavailable: {exc}")

    status["error"] = "; ".join(errors) if errors else None
    return status


def _nested_ai_status(health):
    if isinstance(health, dict) and isinstance(health.get("ai"), dict):
        return health["ai"]
    return health if isinstance(health, dict) else {}

def detect_ai_status() -> Dict:
    checked_at = datetime.utcnow().isoformat(timespec="seconds") + "Z"
    status = {
        "service": "unknown",
        "service_online": None,
        "backend": "unknown",
        "cuda_available": None,
        "gpu_name": None,
        "ai_review_available": None,
        "ai_recovery_enabled": ai_recovery_enabled(),
        "checked_at": checked_at,
        "error": None,
        "health": None,
    }

    if ai_engine is None or not hasattr(ai_engine, "health"):
        status.update(
            service="offline",
            service_online=False,
            ai_review_available=False,
            error="AI service health check unavailable",
        )
        return status

    try:
        health = ai_engine.health()
    except Exception as exc:
        status.update(
            service="offline",
            service_online=False,
            ai_review_available=False,
            error=f"AI service health check failed: {exc}",
        )
        return status

    status["health"] = health
    if not isinstance(health, dict):
        status["error"] = "Invalid AI service health response"
        return status

    online = bool(health.get("ok"))
    ai_status = _nested_ai_status(health)
    cuda_available = ai_status.get("cuda_available")
    pair_match_implemented = ai_status.get("pair_match_implemented")
    backend = "cuda" if cuda_available is True else ("cpu" if cuda_available is False else "unknown")

    status.update(
        service="online" if online else "offline",
        service_online=online,
        backend=backend,
        cuda_available=cuda_available if isinstance(cuda_available, bool) else None,
        gpu_name=ai_status.get("gpu_name") or ai_status.get("device"),
        ai_review_available=(online and pair_match_implemented) if isinstance(pair_match_implemented, bool) else None,
        error=health.get("error") or ai_status.get("model_error"),
    )
    return status

def crop_to_4x3(im: Image.Image) -> Image.Image:
    w,h = im.size; target = 4/3; cur = w/h
    if abs(cur-target) < 0.005: return im
    if cur > target:
        new_w = int(h*target); x0 = (w-new_w)//2; return im.crop((x0,0,x0+new_w,h))
    new_h = int(w/target); y0 = (h-new_h)//2; return im.crop((0,y0,w,y0+new_h))

def load_images(folder: Path):
    exts = {".jpg",".jpeg",".png",".bmp",".tif",".tiff",".webp"}
    return storage.list_files(folder, exts)

def write_debug_report(out_path: Path, progress: dict):
    debug_path = out_path.with_suffix(".debug.json")
    debug_data = {
        "report": str(out_path),
        "backend_used": progress.get("backend_used"),
        "backend_requested": progress.get("backend_requested"),
        "processing_time": {
            "matching_time_sec": progress.get("matching_time_sec", 0),
            "avg_time_per_image_sec": progress.get("avg_time_per_image_sec", 0),
            "images_per_sec": progress.get("images_per_sec", 0),
        },
        "performance": {
            "candidate_comparisons": progress.get("candidate_comparisons", 0),
            "cache_hits": progress.get("cache_hits", 0),
            "feature_cache_hits": progress.get("feature_cache_hits", 0),
            "feature_cache_misses": progress.get("feature_cache_misses", 0),
            "feature_cache_path": progress.get("feature_cache_path"),
        },
        "summary": {
            "before": progress.get("before", 0),
            "after": progress.get("after", 0),
            "matched": progress.get("matched", 0),
            "fallback": progress.get("unmatched", 0),
            "average_confidence": progress.get("average_confidence", 0),
            "lowest_confidence": progress.get("lowest_confidence", 0),
            "highest_confidence": progress.get("highest_confidence", 0),
        },
        "matches": progress.get("match_debug", []),
        "unmatched": progress.get("unmatched_debug", []),
        "image_errors": progress.get("image_errors", []),
        "ai_service_health": progress.get("ai_service_health"),
        "ai_review": {
            "enabled": progress.get("ai_enabled", False),
            "state": progress.get("ai_state", "disabled"),
            "error": progress.get("ai_error"),
            "total": progress.get("ai_total", 0),
            "done": progress.get("ai_done", 0),
            "summary": progress.get("ai_summary") or summarize_ai_results(progress.get("ai_results", [])),
            "results": progress.get("ai_results", []),
            "calibration_error": progress.get("ai_calibration_error"),
        },
        "ai_recovery": progress.get("ai_recovery", {
            "enabled": False,
            "threshold": ai_recovery_threshold(),
            "attempted": 0,
            "recovered": 0,
            "failed": 0,
            "comparisons": 0,
            "results": [],
            "failed_candidates": [],
            "errors": [],
        }),
    }
    with open(debug_path, "w", encoding="utf-8") as fh:
        json.dump(debug_data, fh, indent=2)
    return debug_path

def build_report(input_root: Path, out_path: Path, company: str, zone: str, title: str,
                 threshold: float, progress: dict, should_cancel=None, backend="cpu", ai_review=False, ai_deep_recovery=False):
    requested_backend = str(backend or "cpu").strip().lower() or "cpu"
    main_backend = "cpu" if requested_backend == "auto" else requested_backend
    should_cancel = should_cancel or (lambda: False)
    if should_cancel():
        raise CancelledJob("Cancelled by user")
    before_dir = input_root / "before"
    after_dir  = input_root / "after"
    befores = load_images(before_dir)
    afters  = load_images(after_dir)

    progress.update(total=len(befores), state="preprocess", done=0,
                    before=len(befores), after=len(afters))

    if not befores or not afters:
        raise RuntimeError("No images in before/ or after/")

    pairs = match_pairs(
        befores=befores,
        afters=afters,
        threshold=threshold,
        progress=progress,
        should_cancel=should_cancel,
        cancelled_exception=CancelledJob,
        backend=main_backend,
    )
    progress["backend_requested"] = requested_backend
    progress["main_backend_selected"] = progress.get("backend_used", main_backend)
    if requested_backend == "auto":
        progress["backend_note"] = "Auto selected CPU for main matcher"

    progress["ai_enabled"] = bool(ai_review)
    progress["ai_results"] = []
    progress["ai_error"] = None
    if ai_review or ai_recovery_enabled():
        check_ai_service(progress)
    ai_backend_status = ai_recovery_backend_status(progress) if ai_recovery_enabled() else {
        "backend": "disabled",
        "cuda_available": None,
        "gpu_name": None,
        "fallback_reason": "AI Recovery disabled",
    }
    progress["ai_recovery_backend_decision"] = ai_backend_status
    print(
        "[backend_auto] requested_backend=%s selected_main_backend=%s selected_ai_backend=%s cuda_available=%s gpu=%s fallback_reason=%s" % (
            requested_backend,
            progress.get("main_backend_selected", main_backend),
            ai_backend_status.get("backend"),
            ai_backend_status.get("cuda_available"),
            ai_backend_status.get("gpu_name"),
            ai_backend_status.get("fallback_reason"),
        ),
        flush=True,
    )
    if ai_review:
        run_ai_pair_review(pairs, progress, should_cancel)
        try:
            append_ai_calibration_rows(out_path.name, progress.get("ai_results", []))
        except Exception as exc:
            progress["ai_calibration_error"] = str(exc)
    else:
        progress["ai_state"] = "disabled"

    pairs = run_ai_recovery(befores, afters, pairs, progress, should_cancel, ai_deep_recovery)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.left_margin = Cm(3.0)
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
    normal = doc.styles["Normal"].paragraph_format
    normal.space_before = Pt(0)
    normal.space_after = Pt(0)
    normal.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def add_header(company, zone):
        p = doc.add_paragraph(); r=p.add_run(company.upper()); r.font.size=Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph()
        r2 = p2.add_run(zone.upper())
        r2.bold = True
        r2.underline = True
        r2.font.size = Pt(12)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    def add_title(title):
        p = doc.add_paragraph(); r=p.add_run(title); r.bold=True; r.font.size=Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_pair_table(before_path, after_path):
        from docx.oxml.shared import OxmlElement, qn
        tbl = doc.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in tbl.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tbl_pr = tbl._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); borders.append(el)
        tbl_pr.append(borders)
        try:
            im = Image.open(before_path).convert("RGB"); im = crop_to_4x3(im)
            tmp = REPORT_ROOT / f"~tmp_b.jpg"; im.save(tmp, quality=90)
            cell = tbl.cell(0,0).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run().add_picture(str(tmp), height=Cm(IMG_H_CM)); storage.unlink(tmp)
        except Exception:
            tbl.cell(0,0).text="(gambar gagal)"
        try:
            im = Image.open(after_path).convert("RGB"); im = crop_to_4x3(im)
            tmp = REPORT_ROOT / f"~tmp_a.jpg"; im.save(tmp, quality=90)
            cell = tbl.cell(0,1).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.add_run().add_picture(str(tmp), height=Cm(IMG_H_CM)); storage.unlink(tmp)
        except Exception:
            tbl.cell(0,1).text="(gambar gagal)"
        pL = tbl.cell(1,0).paragraphs[0]; pL.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rL = pL.add_run("Sebelum"); rL.bold=True; rL.font.size=Pt(11)
        pR = tbl.cell(1,1).paragraphs[0]; pR.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rR = pR.add_run("Selepas"); rR.bold=True; rR.font.size=Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_header(company, zone)
    add_title(title)
    per_page = 4
    for i in range(min(per_page, len(pairs))):
        if should_cancel():
            raise CancelledJob("Cancelled by user")
        add_pair_table(pairs[i][0], pairs[i][1])
    if len(pairs) > per_page:
        idx = per_page; pages = 1
        while idx < len(pairs):
            if should_cancel():
                raise CancelledJob("Cancelled by user")
            doc.add_page_break()
            add_header(company, zone)
            for k in range(per_page):
                j = idx + k
                if j >= len(pairs): break
                if should_cancel():
                    raise CancelledJob("Cancelled by user")
                add_pair_table(pairs[j][0], pairs[j][1])
            idx += per_page; pages += 1
            progress["pages"] = pages

    if should_cancel():
        raise CancelledJob("Cancelled by user")
    doc.save(out_path)
    debug_path = write_debug_report(out_path, progress)
    progress.update(state="done", download=str(out_path), debug_download=str(debug_path))

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET_KEY") or os.getenv("SECRET_KEY") or secrets.token_hex(32)
app.permanent_session_lifetime = timedelta(days=30)

ADMIN_USERNAME = os.getenv("ADMIN_USERNAME") or os.getenv("ADMIN_USER")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD") or os.getenv("ADMIN_PASS")
INTERNAL_API_TOKEN = os.getenv("INTERNAL_API_TOKEN")

def admin_auth_enabled():
    return bool(ADMIN_USERNAME and ADMIN_PASSWORD)

def is_admin_logged_in():
    return bool(session.get("admin_authenticated"))

def valid_internal_api_token():
    if not INTERNAL_API_TOKEN:
        return False
    token = request.headers.get("X-Internal-Api-Token", "")
    return bool(token) and secrets.compare_digest(token, INTERNAL_API_TOKEN)

def safe_next_url(value):
    if not value:
        return url_for("index")
    parsed = parse.urlparse(value)
    if parsed.scheme or parsed.netloc or not value.startswith("/"):
        return url_for("index")
    return value

def admin_required(view):
    @wraps(view)
    def wrapped(*args, **kwargs):
        if not admin_auth_enabled() or is_admin_logged_in() or valid_internal_api_token():
            return view(*args, **kwargs)
        return redirect(url_for("login", next=request.full_path.rstrip("?")))
    return wrapped

def no_cache_html(response):
    response = make_response(response)
    response.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

job_manager = JobManager()

def recover_interrupted_report_jobs():
    interrupted_statuses = ("running", "queued", "processing")
    message = "Interrupted by server restart"
    with connection() as conn:
        cursor = conn.execute(
            "UPDATE jobs SET status = ?, error = ?, finished_at = CURRENT_TIMESTAMP WHERE status IN (?, ?, ?)",
            ("error", message, *interrupted_statuses),
        )
        recovered = cursor.rowcount if cursor.rowcount is not None else 0
        conn.commit()
    logger.info("Recovered %s interrupted report job(s) at startup", recovered)
    return recovered

recover_interrupted_report_jobs()


def max_concurrent_report_jobs():
    try:
        return max(1, int(os.getenv("MAX_CONCURRENT_REPORT_JOBS", "1")))
    except ValueError:
        return 1


class BoundedReportQueue:
    def __init__(self, max_workers: int):
        self.max_workers = max(1, int(max_workers or 1))
        self._executor = ThreadPoolExecutor(max_workers=self.max_workers, thread_name_prefix="report-job")
        logger.info("Report job queue initialized with max_workers=%s", self.max_workers)

    def submit(self, target, *args, **kwargs):
        future = self._executor.submit(target, *args, **kwargs)
        logger.info("Report job queued for bounded execution")
        return future


queue_manager = BoundedReportQueue(max_concurrent_report_jobs())
worker_upload_manager = WorkerUploadManager(storage)
REQUIRED_JOB_PAYLOAD_FIELDS = ("month", "site", "task", "company", "zone", "title", "threshold")
RETRYABLE_STATUSES = {"error", "failed", "cancelled"}
RERUN_STATUSES = {"done"}

def sse_stream(job_id):
    q = job_manager.get_events(job_id)
    yield f"event: ping\ndata: keepalive\n\n"
    state = job_manager.get_state(job_id)
    if state:
        yield f"data: {json.dumps(dict(state))}\n\n"
    while True:
        msg = q.get()
        yield f"data: {msg}\n\n"

def push(job_id):
    state = job_manager.get_state(job_id)
    if state:
        job_manager.publish(job_id, dict(state))

def create_report_job(payload):
    month = str(payload["month"]).strip()
    site = str(payload["site"]).strip()
    task = str(payload["task"]).strip()
    company = str(payload["company"]).strip()
    zone = str(payload["zone"]).strip()
    title = str(payload["title"]).strip()
    threshold = float(payload["threshold"])
    backend = str(payload.get("backend", "cpu")).strip().lower() or "cpu"
    ai_review = str(payload.get("ai_review", "")).lower() in {"1", "true", "yes", "on"}
    ai_deep_recovery = bool_from_value(payload.get("ai_deep_recovery"))
    job_id = uuid.uuid4().hex
    job_payload = dict(payload)
    job_payload["threshold"] = threshold
    job_payload["backend"] = backend
    job_payload["ai_review"] = ai_review
    job_payload["ai_deep_recovery"] = ai_deep_recovery
    job_manager.create(
        job_id,
        name=f"{month} {site} {task}",
        payload=job_payload,
        state={"state":"queued","done":0,"total":0,"matched":0,"unmatched":0,"before":0,"after":0,"pages":0,"ai_enabled":ai_review,"ai_results":[],"ai_error":None,"ai_state":"disabled"},
    )
    logger.info("Report job %s queued: %s %s %s", job_id, month, site, task)
    queue_manager.submit(run_queued_report_job, job_id, month, site, task, company, zone, title, threshold, backend, ai_review, ai_deep_recovery)
    return job_id

def clean_saved_payload(job):
    payload = dict((job or {}).get("payload") or {})
    payload.pop("_state", None)
    return payload

def validate_report_payload(payload):
    missing = [field for field in REQUIRED_JOB_PAYLOAD_FIELDS if payload.get(field) in (None, "")]
    if missing:
        return f"Missing saved job payload fields: {', '.join(missing)}"
    try:
        float(payload["threshold"])
    except (TypeError, ValueError):
        return "Invalid saved job threshold"
    return None

def run_queued_report_job(job_id, month, site, task, company, zone, title, threshold, backend="cpu", ai_review=False, ai_deep_recovery=False):
    logger.info("Report job %s starting from queue", job_id)
    try:
        run_job(job_id, month, site, task, company, zone, title, threshold, backend, ai_review, ai_deep_recovery)
    finally:
        job = job_manager.get_state(job_id)
        status = (job or {}).get("state", "unknown")
        logger.info("Report job %s finished with status=%s", job_id, status)


def run_job(job_id, month, site, task, company, zone, title, threshold, backend="cpu", ai_review=False, ai_deep_recovery=False):
    try:
        job = job_manager.get_state(job_id)
        if job_manager.is_cancelled(job_id):
            raise CancelledJob("Cancelled by user")
        job.update(state="starting", done=0, total=0)
        logger.info("Report job %s running: %s %s %s", job_id, month, site, task)
        notify_telegram(f"Report job started: {month} {site} {task}")
        push(job_id)
        input_root = DATA_ROOT / month / site / task
        legacy_name = "drainage" if task == "drainage_cleaning" else task
        out_name = f"{month}_{site}_{legacy_name}.docx"
        out_path = REPORT_ROOT / out_name
        build_report(input_root, out_path, company, zone, title, threshold, job, lambda: job_manager.is_cancelled(job_id), backend, ai_review, ai_deep_recovery)
        logger.info("Report job %s completed: %s", job_id, out_name)
        notify_telegram(f"Report job completed: {out_name}")
        push(job_id)
    except CancelledJob:
        job = job_manager.get_state(job_id)
        if job:
            job.update(state="cancelled", error="Cancelled by user", download=None)
        logger.info("Report job %s cancelled", job_id)
        push(job_id)
    except Exception as e:
        job = job_manager.get_state(job_id)
        if job:
            job.update(state="error", error=str(e))
        logger.exception("Report job %s failed", job_id)
        notify_telegram(f"Report job failed: {month} {site} {task}\n{e}")
        push(job_id)

@app.route("/")
@admin_required
def index():
    ms = scan_months_sites()
    return no_cache_html(render_template("index.html", months_sites=ms, tasks=report_tasks(), defaults=report_defaults()))

@app.route("/login", methods=["GET", "POST"])
def login():
    if not admin_auth_enabled():
        return redirect(url_for("index"))

    next_url = safe_next_url(request.values.get("next"))
    error = None

    if request.method == "POST":
        username = request.form.get("username", "")
        password = request.form.get("password", "")
        if secrets.compare_digest(username, ADMIN_USERNAME) and secrets.compare_digest(password, ADMIN_PASSWORD):
            session.clear()
            session.permanent = request.form.get("remember") == "1"
            session["admin_authenticated"] = True
            return redirect(next_url)
        error = "Invalid username or password"

    return no_cache_html(render_template("login.html", error=error, next_url=next_url))

@app.get("/logout")
def logout():
    session.clear()
    return redirect(url_for("login") if admin_auth_enabled() else url_for("index"))


def current_month_value():
    return datetime.now().strftime("%Y-%m")

WORKER_ALLOWED_IMAGE_EXTENSIONS = {"jpg", "jpeg", "png", "webp", "heic", "heif"}

def worker_upload_limit(name: str, default: int) -> int:
    try:
        return max(1, int(os.getenv(name, str(default))))
    except ValueError:
        return default

def worker_upload_file_size(upload) -> int:
    stream = upload.stream
    position = stream.tell()
    stream.seek(0, os.SEEK_END)
    size = stream.tell()
    stream.seek(position)
    return size

def validate_worker_upload_payload(data):
    site = str(data.get("site", "")).strip()
    task = str(data.get("task", "")).strip()
    month = str(data.get("month", "")).strip()
    worker_name = str(data.get("worker_name", "")).strip()
    if not re.match(r"^\d{4}-\d{2}$", month):
        return None, "Month must use YYYY-MM"
    site_row = project_manager.find_site(site)
    if not site_row:
        return None, "Unknown site"
    task_row = project_manager.task_by_slug(task)
    if not task_row:
        return None, "Unknown task"
    return {"site": site_row["name"], "task": task_row["slug"], "month": month, "worker_name": worker_name}, None

@app.get("/worker-upload")
def worker_upload():
    return no_cache_html(render_template(
        "worker_upload.html",
        sites=project_manager.list_sites(),
        tasks=report_tasks(),
        current_month=current_month_value(),
    ))

@app.get("/api/worker/uploads")
def worker_uploads_list():
    limit = request.args.get("limit", "50")
    status = request.args.get("status", "").strip().lower()
    try:
        jobs = worker_upload_manager.list_jobs(int(limit), status=status)
    except ValueError:
        jobs = worker_upload_manager.list_jobs(status=status)
    return jsonify({"ok": True, "uploads": jobs})

@app.post("/api/worker/uploads")
def worker_uploads_create():
    data = request.get_json(silent=True) or request.form.to_dict()
    payload, error = validate_worker_upload_payload(data)
    if error:
        return jsonify({"ok": False, "error": error}), 400
    try:
        job = worker_upload_manager.create(**payload)
        return jsonify({"ok": True, "upload": job})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.get("/api/worker/uploads/<upload_id>")
def worker_uploads_get(upload_id):
    job = worker_upload_manager.get(upload_id)
    if not job:
        return jsonify({"ok": False, "error": "Upload job not found"}), 404
    return jsonify({"ok": True, "upload": job})

@app.post("/api/worker/uploads/<upload_id>/files")
def worker_uploads_add_files(upload_id):
    when = request.form.get("when", "").strip().lower()
    uploads = [item for item in request.files.getlist("photos") if item and item.filename]
    if not uploads:
        return jsonify({"ok": False, "error": "No photos uploaded"}), 400

    max_files_per_request = worker_upload_limit("WORKER_MAX_FILES_PER_REQUEST", 50)
    if len(uploads) > max_files_per_request:
        return jsonify({"ok": False, "error": f"Maximum {max_files_per_request} files per upload request"}), 400

    job = worker_upload_manager.get(upload_id)
    if not job:
        return jsonify({"ok": False, "error": "Upload job not found"}), 404

    max_files_per_job = worker_upload_limit("WORKER_MAX_FILES_PER_JOB", 300)
    current_file_count = len(job.get("files") or [])
    if current_file_count + len(uploads) > max_files_per_job:
        return jsonify({"ok": False, "error": f"Maximum {max_files_per_job} files per worker job"}), 400

    max_file_mb = worker_upload_limit("WORKER_MAX_FILE_MB", 20)
    max_file_bytes = max_file_mb * 1024 * 1024
    for item in uploads:
        filename = Path(item.filename or "").name
        extension = Path(filename).suffix.lower().lstrip(".")
        if extension not in WORKER_ALLOWED_IMAGE_EXTENSIONS:
            allowed = ", ".join(sorted(WORKER_ALLOWED_IMAGE_EXTENSIONS))
            return jsonify({"ok": False, "error": f"Only image uploads are allowed ({allowed})"}), 400
        if worker_upload_file_size(item) > max_file_bytes:
            return jsonify({"ok": False, "error": f"{filename} exceeds the {max_file_mb} MB image size limit"}), 400

    added = []
    try:
        for item in uploads:
            added.append(worker_upload_manager.add_file(upload_id, when, item))
        job = worker_upload_manager.get(upload_id)
        return jsonify({"ok": True, "files": added, "upload": job})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.delete("/api/worker/uploads/<upload_id>/files/<file_id>")
def worker_uploads_delete_file(upload_id, file_id):
    try:
        removed = worker_upload_manager.delete_file(upload_id, file_id)
        job = worker_upload_manager.get(upload_id)
        return jsonify({"ok": True, "removed": removed, "upload": job})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.post("/api/worker/uploads/<upload_id>/ready")
def worker_uploads_ready(upload_id):
    try:
        job = worker_upload_manager.mark_ready(upload_id)
        return jsonify({"ok": True, "upload": job})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 400

@app.get("/worker-upload/file/<upload_id>/<file_id>")
def worker_upload_file(upload_id, file_id):
    job = worker_upload_manager.get(upload_id)
    if not job:
        return "Not found", 404
    item = next((entry for entry in job.get("files", []) if entry.get("id") == file_id), None)
    if not item:
        return "Not found", 404
    variant = request.args.get("variant", "thumb").strip().lower()
    path_key = "path" if variant == "preview" else "thumbnail_path"
    path = Path(item.get(path_key) or item.get("path") or "").resolve()
    try:
        path.relative_to(DATA_ROOT.resolve())
    except ValueError:
        return "Not found", 404
    if not path.exists() or not path.is_file():
        return "Not found", 404
    return send_file(path)

@app.get("/settings")
@admin_required
def settings():
    return no_cache_html(render_template(
        "settings.html",
        companies=project_manager.list_companies(),
        projects=project_manager.list_projects(),
        sites=project_manager.list_sites(),
        categories=project_manager.list_categories(),
        tasks=project_manager.list_tasks(),
    ))

@app.get("/api/settings/<entity>")
@admin_required
def settings_list(entity):
    try:
        mapping = {
            "companies": project_manager.list_companies,
            "projects": project_manager.list_projects,
            "sites": project_manager.list_sites,
            "categories": project_manager.list_categories,
            "tasks": project_manager.list_tasks,
        }
        return jsonify(ok=True, items=mapping[entity]())
    except KeyError:
        return jsonify(ok=False, error="Unknown settings entity"), 404

@app.post("/api/settings/<entity>")
@admin_required
def settings_create(entity):
    data = request.get_json(silent=True) or request.form.to_dict()
    try:
        return jsonify(ok=True, item=project_manager.create(entity, data))
    except Exception as exc:
        return jsonify(ok=False, error=str(exc)), 400

@app.put("/api/settings/<entity>/<item_id>")
@admin_required
def settings_update(entity, item_id):
    data = request.get_json(silent=True) or request.form.to_dict()
    try:
        return jsonify(ok=True, item=project_manager.update(entity, item_id, data))
    except Exception as exc:
        return jsonify(ok=False, error=str(exc)), 400

@app.delete("/api/settings/<entity>/<item_id>")
@admin_required
def settings_delete(entity, item_id):
    try:
        project_manager.delete(entity, item_id)
        return jsonify(ok=True)
    except Exception as exc:
        return jsonify(ok=False, error=str(exc)), 400

@app.post("/start")
@admin_required
def start():
    data = request.form
    month = data.get("month","").strip()
    site  = data.get("site","").strip()
    defaults = report_defaults()
    task  = data.get("task", defaults["task"]).strip()
    company = data.get("company", defaults["company"]).strip()
    zone    = data.get("zone", f"{site} ZONE").strip()
    task_row = project_manager.task_by_slug(task)
    title   = data.get("title", (task_row or {}).get("title") or defaults["title"]).strip()
    threshold = float(data.get("threshold", defaults["threshold"]))
    backend = data.get("backend", "cpu").strip().lower() or "cpu"
    ai_review = data.get("ai_review") == "1"
    ai_deep_recovery = data.get("ai_deep_recovery") == "1"
    if not month or not site:
        return jsonify({"ok": False, "error": "Please choose a month and a site"}), 400
    payload = {"month": month, "site": site, "task": task, "company": company, "zone": zone, "title": title, "threshold": threshold, "backend": backend, "ai_review": ai_review, "ai_deep_recovery": ai_deep_recovery}
    job_id = create_report_job(payload)
    return jsonify({"ok": True, "job_id": job_id})

@app.get("/progress/<job_id>")
def progress(job_id):
    if not job_manager.get_events(job_id):
        return "no such job", 404
    return Response(sse_stream(job_id), mimetype="text/event-stream")

@app.get("/api/jobs")
@admin_required
def jobs_list():
    limit = request.args.get("limit", "50")
    status = request.args.get("status", "").strip()
    search = request.args.get("search", "").strip()
    try:
        jobs = job_manager.list_jobs(int(limit), status=status, search=search)
    except ValueError:
        jobs = job_manager.list_jobs(status=status, search=search)
    return jsonify({"ok": True, "jobs": jobs, "counts": job_manager.job_counts()})

@app.get("/api/jobs/<job_id>")
@admin_required
def jobs_get(job_id):
    job = job_manager.get_job(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404
    return jsonify({"ok": True, "job": job})

@app.post("/api/jobs/<job_id>/retry")
@admin_required
def jobs_retry(job_id):
    job = job_manager.get_job(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404

    status = (job.get("status") or "").lower()
    if status not in RETRYABLE_STATUSES and status not in RERUN_STATUSES:
        return jsonify({"ok": False, "error": "Only failed, cancelled, or done jobs can be retried"}), 400

    payload = clean_saved_payload(job)
    error = validate_report_payload(payload)
    if error:
        return jsonify({"ok": False, "error": error}), 400

    mode = "duplicate" if status in RERUN_STATUSES else "retry"
    payload["duplicate_of" if mode == "duplicate" else "retry_of"] = job_id
    new_job_id = create_report_job(payload)
    return jsonify({"ok": True, "job_id": new_job_id, "mode": mode})

@app.post("/api/jobs/<job_id>/cancel")
@admin_required
def jobs_cancel(job_id):
    job = job_manager.get_job(job_id)
    if not job:
        return jsonify({"ok": False, "error": "Job not found"}), 404
    status = (job.get("status") or "").lower()
    if status not in {"queued", "starting", "preprocess", "matching", "running"}:
        return jsonify({"ok": False, "error": "Only running or queued jobs can be cancelled"}), 400
    job_manager.cancel(job_id)
    notify_telegram(f"Report job cancelled: {job.get('name') or job_id}")
    return jsonify({"ok": True, "job_id": job_id, "status": "cancelled"})

@app.get("/api/folders/scan")
@admin_required
def folders_scan():
    months_sites = scan_months_sites()
    return jsonify(
        {
            "ok": True,
            "months_sites": months_sites,
            "counts": {
                "months": len(months_sites),
                "sites": sum(len(sites) for sites in months_sites.values()),
            },
        }
    )

@app.get("/api/system/gpu")
def system_gpu():
    return jsonify({"ok": True, "gpu": detect_gpu_status()})

def health_database_check():
    with connection() as conn:
        conn.execute("SELECT 1").fetchone()


def health_writable_check(path: Path):
    probe = path / f".healthcheck-{uuid.uuid4().hex}"
    path.mkdir(parents=True, exist_ok=True)
    probe.write_text("ok")
    probe.unlink(missing_ok=True)

@app.get("/health")
def health():
    checks = {}

    def record(name, callback):
        try:
            callback()
            checks[name] = {"ok": True}
        except Exception as exc:
            checks[name] = {"ok": False, "error": str(exc)}

    record("app", lambda: True)
    record("database", health_database_check)
    record("reports_writable", lambda: health_writable_check(REPORT_ROOT))
    record("photos_writable", lambda: health_writable_check(DATA_ROOT))

    worker_cache = getattr(worker_upload_manager, "cache_root", None)
    if worker_cache:
        record("worker_upload_cache_writable", lambda: health_writable_check(Path(worker_cache)))

    ok = all(item.get("ok") for item in checks.values())
    return jsonify({"ok": ok, "checks": checks}), 200 if ok else 503

@app.get("/api/system/ai")
def system_ai():
    return jsonify({"ok": True, "ai": detect_ai_status()})

@app.post("/api/vision/analyze")
@admin_required
def vision_analyze():
    data = request.get_json(silent=True) or request.form.to_dict()
    month = str(data.get("month", "")).strip()
    site = str(data.get("site", "")).strip()
    task = str(data.get("task", "")).strip()
    try:
        max_distance = max(1, min(int(data.get("max_distance", 5) or 5), 10))
    except (TypeError, ValueError):
        max_distance = 5
    if not month or not site or not task:
        return jsonify({"ok": False, "error": "month, site, and task are required"}), 400

    root = DATA_ROOT / month / site / task
    try:
        root.resolve().relative_to(DATA_ROOT.resolve())
    except ValueError:
        return jsonify({"ok": False, "error": "Invalid folder path"}), 400
    if not root.exists():
        return jsonify({"ok": False, "error": "Photo folder not found"}), 404

    result = analyze_folder(root, cache=VisionCache(), max_distance=max_distance)
    return jsonify({"ok": True, **result})

@app.get("/download")
@admin_required
def download():
    path_value = request.args.get("path", "").strip()
    if not path_value:
        return "Not found", 404

    requested = Path(path_value)
    if requested.is_absolute() or ".." in requested.parts:
        return "Not found", 404

    report_root = REPORT_ROOT.resolve()
    path = report_root / requested

    try:
        for parent in path.parents:
            if parent == report_root.parent:
                break
            if parent.exists() and parent.is_symlink():
                return "Not found", 404
        if path.exists() and path.is_symlink():
            return "Not found", 404
        resolved = path.resolve(strict=True)
        resolved.relative_to(report_root)
    except (FileNotFoundError, RuntimeError, ValueError):
        return "Not found", 404

    if not resolved.is_file():
        return "Not found", 404
    return send_file(resolved, as_attachment=True)

@app.get("/download/job/<job_id>")
@admin_required
def download_job(job_id):
    job = job_manager.get_job(job_id)
    if not job or not job.get("result_path"):
        return "Not found", 404
    path = Path(job["result_path"]).resolve()
    try:
        path.relative_to(REPORT_ROOT.resolve())
    except ValueError:
        return "Not found", 404
    if not path.exists() or not path.is_file():
        return "Not found", 404
    return send_file(path, as_attachment=True)

@app.get("/download/debug/<job_id>")
@admin_required
def download_debug(job_id):
    job = job_manager.get_job(job_id)
    if not job or not job.get("result_path"):
        return "Not found", 404
    report_path = Path(job["result_path"]).resolve()
    try:
        report_path.relative_to(REPORT_ROOT.resolve())
    except ValueError:
        return "Not found", 404
    debug_path = report_path.with_suffix(".debug.json")
    if not debug_path.exists() or not debug_path.is_file():
        return "Not found", 404
    return send_file(debug_path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080, debug=False)
