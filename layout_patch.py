
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def _clear_table_borders(tbl):
    tblPr = tbl._tbl.tblPr
    borders = tblPr.tblBorders if tblPr is not None and getattr(tblPr, "tblBorders", None) is not None else OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tblPr.append(borders)

def add_pairs_table_4pp(doc, pairs, image_height_cm: float=4.5, label_lang: str='ms'):
    if not pairs: return
    label_before = "Sebelum" if label_lang=='ms' else "Before"
    label_after  = "Selepas" if label_lang=='ms' else "After"
    per_page = 4
    for idx, pair in enumerate(pairs, start=1):
        tbl = doc.add_table(rows=2, cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        _clear_table_borders(tbl)
        cells_img = tbl.rows[0].cells; cells_lbl = tbl.rows[1].cells
        for c in cells_img + cells_lbl: c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pb = cells_img[0].paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try: pb.add_run().add_picture(str(pair.before), height=Cm(image_height_cm))
        except Exception: pb.add_run(f"[Image error: {pair.before.name}]")
        pa = cells_img[1].paragraphs[0]; pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try: pa.add_run().add_picture(str(pair.after),  height=Cm(image_height_cm))
        except Exception: pa.add_run(f"[Image error: {pair.after.name}]")
        plb = cells_lbl[0].paragraphs[0]; plb.alignment = WD_ALIGN_PARAGRAPH.CENTER; plb.add_run(label_before).bold = True
        pla = cells_lbl[1].paragraphs[0]; pla.alignment = WD_ALIGN_PARAGRAPH.CENTER; pla.add_run(label_after).bold  = True
        if idx % per_page == 0 and idx != len(pairs): doc.add_page_break()
