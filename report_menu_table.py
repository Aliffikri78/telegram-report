from __future__ import annotations
import os, re
from dataclasses import dataclass
from pathlib import Path
from typing import List, Dict

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from PIL import Image, ImageOps
import numpy as np

# ---- ENV FLAGS ----
USE_RESNET    = os.getenv("USE_RESNET", "true").lower() == "true"   # GPU feature matcher
IGNORE_NAMES  = os.getenv("IGNORE_NAMES", "true").lower() == "true" # visual-only pairing

# ---- OPTIONAL DEPS ----
try:
    import cv2
except Exception:
    cv2 = None

try:
    import torch
    TORCH_OK = torch.cuda.is_available()
except Exception:
    torch, TORCH_OK = None, False

# 4-per-page renderer (borders hidden)
try:
    from layout_patch import add_pairs_table_4pp as _render_pairs
except Exception:
    _render_pairs = None

# ---- UTILS ----
IMG_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".tif", ".tiff", ".webp"}

def is_image_file(p: Path) -> bool:
    return p.suffix.lower() in IMG_EXTS

def tnr(run, size_pt: int, bold: bool=False, underline: bool=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Times New Roman")
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.underline = underline

def apply_print_layout(doc, base_margin_cm=2.0, binding_extra_cm=1.0):
    """Tight layout + binding-safe left margin; kill paragraph spacing."""
    for section in doc.sections:
        section.top_margin = Cm(base_margin_cm)
        section.bottom_margin = Cm(base_margin_cm)
        section.right_margin = Cm(base_margin_cm)
        section.left_margin = Cm(base_margin_cm + binding_extra_cm)  # e.g., 3.0 cm
        section.header_distance = Cm(1.0)
        section.footer_distance = Cm(1.0)
    style = doc.styles['Normal']
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

def set_page_header(doc, company: str, font_pt: int = 16):
    """Put company name into the REAL Word header (centered) on all pages."""
    for section in doc.sections:
        hdr = section.header
        p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
        try:
            p.clear()
        except Exception:
            p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(company)
        tnr(run, font_pt, bold=False, underline=False)
        section.header_distance = Cm(1.0)

@dataclass
class Pair:
    before: Path
    after: Path
    score: float

def normalize_name(p: Path) -> str:
    name = p.stem.lower()
    name = re.sub(r'(before|after|sebelum|selepas|edited|final|ori|original)', '', name)
    name = re.sub(r'[^a-z0-9]+', ' ', name).strip()
    name = re.sub(r'\s+', ' ', name)
    return name

def filename_similarity(a: str, b: str) -> float:
    sa, sb = set(a.split()), set(b.split())
    if not sa or not sb: return 0.0
    inter, union = len(sa & sb), len(sa | sb)
    return inter/union if union else 0.0

# ---- SIMILARITY (GPU first, then CPU fallback) ----
_model_cache = {"model": None, "tfm": None, "device": "cpu"}

def _gpu_feature_sim(imgA: Image.Image, imgB: Image.Image) -> float:
    if not (USE_RESNET and torch is not None and TORCH_OK):
        return -1.0
    try:
        import torchvision.transforms as T
        from torchvision.models import resnet18, ResNet18_Weights
        if _model_cache["model"] is None:
            device = "cuda"
            m = resnet18(weights=ResNet18_Weights.DEFAULT).to(device).eval()
            tfm = T.Compose([
                T.Resize((224,224)), T.ToTensor(),
                T.Normalize(mean=[0.485,0.456,0.406], std=[0.229,0.224,0.225])
            ])
            _model_cache.update(model=m, tfm=tfm, device=device)
            print("[report] Using GPU ResNet matcher", flush=True)
        m = _model_cache["model"]; tfm = _model_cache["tfm"]; device = _model_cache["device"]
        with torch.inference_mode():
            a = tfm(imgA.convert("RGB")).unsqueeze(0).to(device)
            b = tfm(imgB.convert("RGB")).unsqueeze(0).to(device)
            # pooled features (works for torchvision<=0.18)
            feats_a = m.forward_features(a) if hasattr(m,"forward_features") else m.avgpool(m.layer4(m.layer3(m.layer2(m.layer1(m.relu(m.bn1(m.conv1(a)))))))); feats_a = torch.flatten(feats_a,1)
            feats_b = m.forward_features(b) if hasattr(m,"forward_features") else m.avgpool(m.layer4(m.layer3(m.layer2(m.layer1(m.relu(m.bn1(m.conv1(b)))))))); feats_b = torch.flatten(feats_b,1)
            fa = torch.nn.functional.normalize(feats_a, dim=1)
            fb = torch.nn.functional.normalize(feats_b, dim=1)
            sim = torch.clamp((fa*fb).sum(dim=1), 0, 1).item()
            return float(sim)
    except Exception as e:
        print(f"[report] GPU feature error: {e}", flush=True)
        return -1.0

def _cpu_hist_corr(imgA: Image.Image, imgB: Image.Image) -> float:
    try:
        a = np.array(imgA.convert("RGB"))
        b = np.array(imgB.convert("RGB"))
        if cv2 is not None:
            ah = cv2.cvtColor(a, cv2.COLOR_RGB2HSV)
            bh = cv2.cvtColor(b, cv2.COLOR_RGB2HSV)
            a_hist = cv2.calcHist([ah],[0,1,2],None,[32,32,32],[0,180,0,256,0,256])
            b_hist = cv2.calcHist([bh],[0,1,2],None,[32,32,32],[0,180,0,256,0,256])
            cv2.normalize(a_hist,a_hist); cv2.normalize(b_hist,b_hist)
            corr = cv2.compareHist(a_hist,b_hist,cv2.HISTCMP_CORREL)
            return float((corr+1.0)/2.0)
        # grayscale fallback
        a_gray = np.array(imgA.convert("L"))
        b_gray = np.array(imgB.convert("L"))
        a_hist,_ = np.histogram(a_gray,bins=64,range=(0,255),density=True)
        b_hist,_ = np.histogram(b_gray,bins=64,range=(0,255),density=True)
        num = np.sum((a_hist-a_hist.mean())*(b_hist-b_hist.mean()))
        den = np.sqrt(np.sum((a_hist-a_hist.mean())**2)*np.sum((b_hist-b_hist.mean())**2))+1e-8
        return float((num/den+1.0)/2.0)
    except Exception:
        return 0.0

def visual_similarity(imgA: Image.Image, imgB: Image.Image) -> float:
    g = _gpu_feature_sim(imgA, imgB)
    if g >= 0: return g
    if _model_cache["model"] is None:
        print("[report] Using CPU matcher", flush=True)
    return _cpu_hist_corr(imgA, imgB)

# ---- CORE PIPELINE ----
def load_pil(p: Path) -> Image.Image:
    img = Image.open(p)
    try: img = ImageOps.exif_transpose(img)
    except Exception: pass
    return img

def find_pairs(before_dir: Path, after_dir: Path, threshold: float=0.7, visual_weight: float=0.5) -> List[Pair]:
    befores = [p for p in before_dir.glob("**/*") if p.is_file() and is_image_file(p)]
    afters  = [p for p in after_dir.glob("**/*")  if p.is_file() and is_image_file(p)]

    used_after = set()
    pairs: List[Pair] = []
    cache: Dict[Path, str] = {}

    for b in sorted(befores):
        nb = cache.setdefault(b, normalize_name(b))
        try: img_b = load_pil(b)
        except Exception: img_b = None

        best_score, best_a = -1.0, None
        for a in afters:
            if a in used_after: continue
            na = cache.setdefault(a, normalize_name(a))
            name_sim = 0.0 if IGNORE_NAMES else filename_similarity(nb, na)
            visual_corr = visual_similarity(img_b, load_pil(a)) if img_b else 0.0
            score = visual_corr if IGNORE_NAMES else (1.0 - visual_weight) * name_sim + visual_weight * visual_corr
            if score > best_score:
                best_score, best_a = score, a

        if best_a is not None and best_score >= threshold:
            pairs.append(Pair(before=b, after=best_a, score=best_score))
            used_after.add(best_a)

    return pairs

def add_zone(doc: Document, zone: str):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(zone.upper()); tnr(run, 12, bold=True, underline=True)

def add_section_title(doc: Document, idx: int, title: str):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{idx}. {title}"); tnr(run, 12, bold=True)

def build_report(photos_root: Path, reports_root: Path, month: str, site: str, task: str,
                 company: str=None, threshold: float=None, visual_weight: float=None,
                 image_height_cm: float=None, label_lang: str=None) -> Path:

    # Defaults from env
    company        = company        or os.getenv("COMPANY", "HW UNGGUL (901587-V)")
    threshold      = threshold      if threshold      is not None else float(os.getenv("THRESHOLD", "0.30"))
    visual_weight  = visual_weight  if visual_weight  is not None else float(os.getenv("VISUAL_WEIGHT", "0.80"))
    image_height_cm= image_height_cm if image_height_cm is not None else float(os.getenv("IMG_H_CM", "4.3"))
    label_lang     = label_lang     or os.getenv("LABEL_LANG", "ms")

    before_dir = photos_root / month / site / task / "before"
    after_dir  = photos_root / month / site / task / "after"
    if not before_dir.exists() or not after_dir.exists():
        raise FileNotFoundError(f"Folders not found: {before_dir} or {after_dir}")

    pairs = find_pairs(before_dir, after_dir, threshold=threshold, visual_weight=visual_weight)

    doc = Document()
    apply_print_layout(doc, base_margin_cm=2.0, binding_extra_cm=1.0)
    set_page_header(doc, company)  # real header
    add_zone(doc, f"{site} ZONE")
    title = "Grass Cutting" if task == "grass_cutting" else "Drainage Cleaning"
    add_section_title(doc, 1 if task == "grass_cutting" else 2, title)

    if _render_pairs:
        _render_pairs(doc, pairs, image_height_cm=image_height_cm, label_lang=label_lang)
    else:
        # simple fallback (shouldn’t trigger if layout_patch.py is present)
        for pair in pairs:
            tbl = doc.add_table(rows=2, cols=2)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cells_img = tbl.rows[0].cells; cells_lbl = tbl.rows[1].cells
            for c in cells_img + cells_lbl: c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            pb = cells_img[0].paragraphs[0]; pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: pb.add_run().add_picture(str(pair.before), height=Cm(image_height_cm))
            except Exception: pb.add_run(f"[Image error: {pair.before.name}]")
            pa = cells_img[1].paragraphs[0]; pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: pa.add_run().add_picture(str(pair.after),  height=Cm(image_height_cm))
            except Exception: pa.add_run(f"[Image error: {pair.after.name}]")
            lbl_b = cells_lbl[0].paragraphs[0]; lbl_b.alignment = WD_ALIGN_PARAGRAPH.CENTER; lbl_b.add_run("Sebelum" if label_lang=='ms' else "Before").bold = True
            lbl_a = cells_lbl[1].paragraphs[0]; lbl_a.alignment = WD_ALIGN_PARAGRAPH.CENTER; lbl_a.add_run("Selepas" if label_lang=='ms' else "After").bold = True

    reports_root.mkdir(parents=True, exist_ok=True)
    outpath = reports_root / f"{month}_{site}_{task}.docx"
    doc.save(outpath)
    return outpath

# ---- OPTIONAL CLI ----
def parse_args(argv=None):
    import argparse
    ap = argparse.ArgumentParser(description="GPU-accelerated report generator")
    ap.add_argument("--photos-root", required=True)
    ap.add_argument("--reports-root", required=True)
    ap.add_argument("--month", required=True)
    ap.add_argument("--site", required=True)
    ap.add_argument("--task", choices=["grass_cutting","drainage_cleaning"], required=True)
    ap.add_argument("--company", default=None)
    ap.add_argument("--threshold", type=float, default=None)
    ap.add_argument("--visual-weight", type=float, default=None)
    ap.add_argument("--image-height-cm", type=float, default=None)
    ap.add_argument("--label-lang", choices=["ms","en"], default=None)
    return ap.parse_args(argv)

def main():
    args = parse_args()
    out = build_report(
        photos_root=Path(args.photos_root),
        reports_root=Path(args.reports_root),
        month=args.month, site=args.site, task=args.task,
        company=args.company, threshold=args.threshold, visual_weight=args.visual_weight,
        image_height_cm=args.image_height_cm, label_lang=args.label_lang
    )
    print(out)

if __name__ == "__main__":
    main()