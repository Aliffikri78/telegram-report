#!/usr/bin/env python3
import os, re, threading, queue, uuid
from pathlib import Path
from typing import List, Dict
from flask import Flask, render_template, request, Response, jsonify, send_file

from PIL import Image
import cv2, numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import imagehash

DATA_ROOT = Path(os.getenv("SAVE_ROOT", "/data/photos")).resolve()
REPORT_ROOT = Path("/data/reports").resolve()
REPORT_ROOT.mkdir(parents=True, exist_ok=True)

def list_dirs(p: Path):
    return [d for d in sorted(p.iterdir()) if d.is_dir()]

def is_month_name(name: str) -> bool:
    import re
    return re.match(r"^\\d{4}-\\d{2}$", name) is not None

def scan_months_sites() -> Dict[str, list]:
    out = {}
    if not DATA_ROOT.exists(): return out
    for m in list_dirs(DATA_ROOT):
        if not is_month_name(m.name): continue
        sites = []
        for s in list_dirs(m):
            if (s / "grass_cutting").exists() or (s / "drainage_cleaning").exists():
                sites.append(s.name)
        if sites:
            out[m.name] = sites
    return out

MAX_SIDE = int(os.getenv("FAST_MAX_SIDE", "1600"))
NFEATURES = int(os.getenv("FAST_NFEATURES", "600"))
TOPK = int(os.getenv("FAST_TOPK", "5"))
RATIO = float(os.getenv("FAST_RATIO", "0.75"))

def crop_to_4x3(im: Image.Image) -> Image.Image:
    w,h = im.size; target = 4/3; cur = w/h
    if abs(cur-target) < 0.005: return im
    if cur > target:
        new_w = int(h*target); x0 = (w-new_w)//2; return im.crop((x0,0,x0+new_w,h))
    new_h = int(w/target); y0 = (h-new_h)//2; return im.crop((0,y0,w,y0+new_h))

def load_images(folder: Path):
    exts = {".jpg",".jpeg",".png",".bmp",".tif",".tiff",".webp"}
    if not folder.exists(): return []
    return [p for p in sorted(folder.iterdir()) if p.suffix.lower() in exts]

def imread_gray_resized(path: Path):
    arr = cv2.imdecode(np.fromfile(str(path), dtype=np.uint8), cv2.IMREAD_GRAYSCALE)
    if arr is None: return None
    h,w = arr.shape; m = max(h,w)
    if m > MAX_SIDE and m>0:
        scale = MAX_SIDE/m
        arr = cv2.resize(arr, (int(w*scale), int(h*scale)), interpolation=cv2.INTER_AREA)
    return arr

def phash(path: Path) -> int:
    with Image.open(path) as im:
        return int(str(imagehash.phash(im.convert("RGB"))), 16)

def hamming(a: int, b: int) -> int:
    return bin(a ^ b).count("1")

def score_pair(des_a, des_b) -> float:
    if des_a is None or des_b is None or len(des_a)==0 or len(des_b)==0: return 0.0
    bf = cv2.BFMatcher(cv2.NORM_HAMMING)
    try:
        matches = bf.knnMatch(des_a, des_b, k=2)
    except cv2.error:
        return 0.0
    good = 0
    for m_n in matches:
        if len(m_n) < 2: continue
        m, n = m_n
        if m.distance < RATIO * n.distance:
            good += 1
    return min(1.0, good/200.0)

def build_report(input_root: Path, out_path: Path, company: str, zone: str, title: str,
                 threshold: float, progress: dict):
    before_dir = input_root / "before"
    after_dir  = input_root / "after"
    befores = load_images(before_dir)
    afters  = load_images(after_dir)

    progress.update(total=len(befores), state="preprocess", done=0,
                    before=len(befores), after=len(afters))

    if not befores or not afters:
        raise RuntimeError("No images in before/ or after/")

    orb = cv2.ORB_create(nfeatures=NFEATURES)
    be_ph, af_ph = [], []
    be_desc, af_desc = {}, {}

    for p in befores:
        be_ph.append(phash(p))
        arr = imread_gray_resized(p)
        k, d = orb.detectAndCompute(arr, None) if arr is not None else (None,None)
        be_desc[p] = d

    for p in afters:
        af_ph.append(phash(p))
        arr = imread_gray_resized(p)
        k, d = orb.detectAndCompute(arr, None) if arr is not None else (None,None)
        af_desc[p] = d

    progress.update(state="matching", done=0, matched=0, unmatched=0)

    pairs = []
    used = {}
    af_items = list(zip(afters, af_ph))
    for idx, b in enumerate(befores):
        dlist = [(j, hamming(be_ph[idx], aph)) for j,(_,aph) in enumerate(af_items) if not used.get(j)]
        dlist.sort(key=lambda x: x[1])
        cand = [j for j,_ in dlist[:max(1, TOPK)]]
        best_j, best_s = None, -1.0
        for j in cand:
            apath = af_items[j][0]
            s = score_pair(be_desc[b], af_desc[apath])
            if s > best_s:
                best_s, best_j = s, j
        if best_j is not None and best_s >= threshold:
            used[best_j] = True
            pairs.append((b, afters[best_j], best_s))
            progress["matched"] += 1
        else:
            j = idx if (idx < len(afters) and not used.get(idx)) else None
            if j is not None:
                used[j] = True
                pairs.append((b, afters[j], -1.0))
                progress["unmatched"] += 1
        progress["done"] = idx + 1

    # DOCX (table layout)
    doc = Document()
    def add_header(company, zone):
        p = doc.add_paragraph(); r=p.add_run(company.upper()); r.bold=True; r.font.size=Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = doc.add_paragraph(zone.upper()); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; p2.runs[0].font.size=Pt(12)
    def add_title(title):
        p = doc.add_paragraph(); r=p.add_run(title); r.bold=True; r.font.size=Pt(12); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    def add_pair_table(before_path, after_path):
        from docx.oxml.shared import OxmlElement, qn
        sec = doc.sections[-1]
        gutter = Inches(0.3)
        max_w = (sec.page_width - sec.left_margin - sec.right_margin - gutter) / 2
        tbl = doc.add_table(rows=2, cols=2)
        # borderless
        tbl_pr = tbl._tbl.tblPr
        borders = OxmlElement('w:tblBorders')
        for edge in ('top','left','bottom','right','insideH','insideV'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); borders.append(el)
        tbl_pr.append(borders)
        # left
        try:
            im = Image.open(before_path).convert("RGB"); im = crop_to_4x3(im)
            tmp = REPORT_ROOT / f"~tmp_b.jpg"; im.save(tmp, quality=90)
            cell = tbl.cell(0,0).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.LEFT
            cell.add_run().add_picture(str(tmp), width=max_w); tmp.unlink(missing_ok=True)
        except Exception:
            tbl.cell(0,0).text="(gambar gagal)"
        # right
        try:
            im = Image.open(after_path).convert("RGB"); im = crop_to_4x3(im)
            tmp = REPORT_ROOT / f"~tmp_a.jpg"; im.save(tmp, quality=90)
            cell = tbl.cell(0,1).paragraphs[0]; cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cell.add_run().add_picture(str(tmp), width=max_w); tmp.unlink(missing_ok=True)
        except Exception:
            tbl.cell(0,1).text="(gambar gagal)"
        # labels
        pL = tbl.cell(1,0).paragraphs[0]; pL.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rL = pL.add_run("SEBELUM"); rL.bold=True; rL.font.size=Pt(11)
        pR = tbl.cell(1,1).paragraphs[0]; pR.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rR = pR.add_run("SELEPAS"); rR.bold=True; rR.font.size=Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_header(os.getenv("COMPANY","HW UNGGUL (901587-V)"), os.getenv("ZONE","ZONE"))
    add_title(os.getenv("TITLE","1. Grass Cutting"))
    per_page = 4
    for i in range(min(per_page, len(pairs))):
        add_pair_table(pairs[i][0], pairs[i][1])
    if len(pairs) > per_page:
        idx = per_page; pages = 1
        while idx < len(pairs):
            doc.add_page_break()
            add_header(os.getenv("COMPANY","HW UNGGUL (901587-V)"), os.getenv("ZONE","ZONE"))
            for k in range(per_page):
                j = idx + k
                if j >= len(pairs): break
                add_pair_table(pairs[j][0], pairs[j][1])
            idx += per_page; pages += 1
            progress["pages"] = pages

    doc.save(out_path)
    progress.update(state="done", download=str(out_path))

app = Flask(__name__)

JOBS = {}
EVENTS = {}

def sse_stream(job_id):
    q = EVENTS[job_id]
    yield f"event: ping\ndata: keepalive\n\n"
    while True:
        msg = q.get()
        yield f"data: {msg}\n\n"

def push(job_id):
    import json
    q = EVENTS.get(job_id)
    if q:
        q.put(json.dumps(JOBS[job_id]))

def run_job(job_id, month, site, task, company, zone, title, threshold):
    try:
        from queue import Queue
        JOBS[job_id].update(state="starting", done=0, total=0)
        push(job_id)
        input_root = DATA_ROOT / month / site / task
        out_name = f"{month}_{site}_{'grass_cutting' if task=='grass_cutting' else 'drainage'}.docx"
        out_path = REPORT_ROOT / out_name
        build_report(input_root, out_path, company, zone, title, threshold, JOBS[job_id])
        push(job_id)
    except Exception as e:
        JOBS[job_id].update(state="error", error=str(e)); push(job_id)

@app.route("/")
def index():
    ms = scan_months_sites()
    return render_template("index.html", months_sites=ms)

@app.post("/start")
def start():
    data = request.form
    month = data.get("month","").strip()
    site  = data.get("site","").strip()
    task  = data.get("task","grass_cutting").strip()
    company = data.get("company","HW UNGGUL (901587-V)").strip()
    zone    = data.get("zone", f"{site} ZONE").strip()
    title   = data.get("title","1. Grass Cutting" if task=="grass_cutting" else "2. Drainage Cleaning").strip()
    threshold = float(data.get("threshold","0.70"))
    if not month or not site:
        return jsonify({"ok": False, "error": "Please choose a month and a site"}), 400
    job_id = uuid.uuid4().hex
    from queue import Queue
    JOBS[job_id] = {"state":"queued","done":0,"total":0,"matched":0,"unmatched":0,"before":0,"after":0,"pages":0}
    EVENTS[job_id] = Queue()
    threading.Thread(target=run_job, args=(job_id, month, site, task, company, zone, title, threshold), daemon=True).start()
    return jsonify({"ok": True, "job_id": job_id})

@app.get("/progress/<job_id>")
def progress(job_id):
    if job_id not in EVENTS:
        return "no such job", 404
    return Response(sse_stream(job_id), mimetype="text/event-stream")

@app.get("/download")
def download():
    path = request.args.get("path")
    if not path or not Path(path).exists():
        return "Not found", 404
    return send_file(path, as_attachment=True)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8080, debug=False)
